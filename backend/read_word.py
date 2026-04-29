"""
Project: Accounts Team
Module: read_word
Author: Dhinakaran Sekar
Email: dhinakaran.s@jubilantenterprises.in
Date: 2026-04-08 11:53:28
"""
import re
from docx import Document

# -------------------------------
# 🔹 COMPILED PATTERNS
# -------------------------------
# Pre-compiling regex expressions is significantly faster for repeated executions
PAN_REGEX = re.compile(r"\(\s*PAN.*", re.IGNORECASE)
PRIMARY_CLIENT_REGEX = re.compile(r"(?:M/s|Mr|Mrs|Ms)\.?\s*(.*?)[,]?$", re.IGNORECASE)
CLIENT_ACCOUNT_REGEX = re.compile(r"From Name:\s*(?:M/s|Mr|Mrs|Ms)\.?\s*(.*?),\s*Bank:", re.IGNORECASE)
LOAN_DATE_REGEX = re.compile(r"Date\s*:\s*(\d{2}-\d{2}-\d{4})")
LOAN_AMOUNT_REGEX = re.compile(r"Rs\.?\s*([\d,]+)\s*/-")


# -------------------------------
# 🔹 COMMON UTILS
# -------------------------------
def clean_text(text):
    # Split and join is faster natively than using re.sub for generic whitespace
    return " ".join(text.split())


# -------------------------------
# 🔹 PRIMARY ACCOUNT
# -------------------------------
def extract_primary_account(paragraphs):
    # Uses the already evaluated list of paragraphs to prevent re-looping the Document
    for i, text in enumerate(paragraphs):
        if "(PAN" in text.upper():
            # ✅ Only consider PAN under "To"
            if i > 0 and paragraphs[i - 1].strip().lower().startswith("to"):
                before_pan = PAN_REGEX.split(text)[0]
                match = PRIMARY_CLIENT_REGEX.search(before_pan.strip())
                if match:
                    return clean_text(match.group(1).rstrip(','))
    return None


# -------------------------------
# 🔹 CLIENT ACCOUNT
# -------------------------------
def extract_client_account(full_text):
    match = CLIENT_ACCOUNT_REGEX.search(full_text)
    return clean_text(match.group(1)) if match else None


# -------------------------------
# 🔹 DATE
# -------------------------------
def extract_loan_date(full_text):
    match = LOAN_DATE_REGEX.search(full_text)
    return match.group(1) if match else None


# -------------------------------
# 🔹 LOAN AMOUNT
# -------------------------------
def extract_loan_amount(full_text):
    match = LOAN_AMOUNT_REGEX.search(full_text)
    return int(match.group(1).replace(",", "")) if match else None


# -------------------------------
# 🔹 TABLE + TOTAL
# -------------------------------
def extract_table_data(doc):
    table_data = []
    total_repayment = None
    
    date_idx = 1
    cheque_idx = 2
    gross_idx = -1
    tds_idx = -1
    net_idx = -1

    for table in doc.tables:
        header_row = None
        for row in table.rows:
            cells = [cell.text.strip().upper() for cell in row.cells]
            if not cells: continue

            # Detect Header Row
            if any(h in cells for h in ["DATE", "CHEQUE NO", "CHEQ NO", "GROSS AMOUNT"]):
                header_row = cells
                for i, h in enumerate(header_row):
                    if "DATE" in h: date_idx = i
                    if "CHEQ" in h: cheque_idx = i
                    if "GROSS" in h: gross_idx = i
                    if "TDS" in h: tds_idx = i
                    if "NET" in h or "TOTAL AMOUNT" in h: net_idx = i
                continue

            # Detect TOTAL row
            if any("TOTAL" in cell for cell in cells):
                try:
                    # Priority: Net/Total column > Sum of Gross+TDS > Last cell
                    val = 0
                    if net_idx != -1:
                        val = float(cells[net_idx].replace(",", ""))
                    elif gross_idx != -1 and tds_idx != -1:
                        val = float(cells[gross_idx].replace(",", "")) + float(cells[tds_idx].replace(",", ""))
                    else:
                        val = float(cells[-1].replace(",", ""))
                    total_repayment = val
                except (ValueError, IndexError):
                    pass
                continue

            # Data rows
            first_cell = cells[0]
            if len(cells) >= 4 and first_cell.isdigit():
                try:
                    amt = 0
                    if net_idx != -1:
                        amt = float(cells[net_idx].replace(",", ""))
                    elif gross_idx != -1 and tds_idx != -1:
                        amt = float(cells[gross_idx].replace(",", "")) + float(cells[tds_idx].replace(",", ""))
                    else:
                        amt = float(cells[3].replace(",", ""))

                    table_data.append({
                        "date": row.cells[date_idx].text.strip(),
                        "cheque_no": row.cells[cheque_idx].text.strip() if cheque_idx < len(cells) else "",
                        "amount": amt,
                        "received_date": None,
                        "payment_date": None
                    })
                except (ValueError, IndexError):
                    continue

    return table_data, total_repayment


# -------------------------------
# 🔹 EXTRACT ALL DATA
# -------------------------------
def extract_word_data(file_path):
    doc = Document(file_path)

    # Cache paragraphs sequence just once and re-use it
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    full_text = "\n".join(paragraphs)

    table_data, total_repayment = extract_table_data(doc)

    return {
        "primary_account": extract_primary_account(paragraphs),
        "client_account": extract_client_account(full_text),
        "loan_date": extract_loan_date(full_text),
        "loan_amount": extract_loan_amount(full_text),
        "table": table_data,
        "total_repayment": total_repayment
    }


# -------------------------------
# 🔹 FINAL OUTPUT BUILDER
# -------------------------------
def build_final_output(data, remaining_accounts):
    loan_amount = data["loan_amount"]

    # 🔸 Validate
    if loan_amount is None:
        raise ValueError("Loan amount not found")

    remaining_total = sum(acc["share"] for acc in remaining_accounts)

    if remaining_total > loan_amount:
        raise ValueError("Remaining shares exceed loan amount")

    total_interest = (data["total_repayment"] - loan_amount) if data["total_repayment"] is not None else 0

    # 🔸 Primary account
    primary_amount = loan_amount - remaining_total
    primary_percentage = round((primary_amount / loan_amount) * 100, 2)
    primary_interest = round((primary_amount / loan_amount) * total_interest, 2)

    # 🔸 Remaining accounts comprehension
    formatted_remaining = [
        {
            "account_name": acc["name"].strip().upper(),
            "share": acc["share"],
            "percentage": round((acc["share"] / loan_amount) * 100, 2),
            "interest_amount": round((acc["share"] / loan_amount) * total_interest, 2),
            "tds": round(((acc["share"] / loan_amount) * total_interest) * 0.10, 2)
        }
        for acc in remaining_accounts
    ]

    # 🔸 Final structure
    return {
        "loan_date": data["loan_date"],
        "loan_amount": loan_amount,
        "total_repayment_amount": data["total_repayment"],
        "total_interest": round(total_interest, 2),

        "primary_account_name": data["primary_account"].strip().upper() if data["primary_account"] else None,
        "primary_account_share": primary_percentage,
        "primary_account_amount": primary_amount,
        "primary_account_interest": primary_interest,

        "client_account_name": data["client_account"],

        "remaining_accounts": formatted_remaining,

        "table": data["table"]
    }

